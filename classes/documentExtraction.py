import docx
import sys
import pandas as pd
from docx.api import Document
import nltk
# nltk.download('punkt')
from nltk.tokenize import sent_tokenize
import re
from tqdm import tqdm


class DocumentModel:
    def __init__(self):
        self=self
    def extractTable(self,item):
        combined_list = []
        each_list = []
        for each_document in tqdm(item):
            document = Document(each_document)
            if document.tables:
                each_list = [[cell.text for cell in row.cells] for row in document.tables[0].rows][1:]
                combined_list.extend(each_list)
                
                # for row in document.tables[0].rows[1:]:
                #     cell_list = []
                #     for cell in row.cells:
                #         cell_list.append(cell.text)
                #     each_list.append(cell_list)
        budget_tabledata = pd.DataFrame(combined_list, columns=["Sl.no", "Department", "Outlay"])
        return (budget_tabledata)



    def extractDescription(self,data):
        data.fillna('nill', inplace=True)
        data['tokenized_sentences'] = data.expense.apply(sent_tokenize)
        data['extracted_sentance']=None
        for index,row in data.iterrows():
            x=''
            value=[]
            for info in row['tokenized_sentences']:
                if bool(re.search(r'([0-9]{4}\-[0-9]{2})',info)) or bool(re.search("₹""\d+", info)):
                    info=str(info)
                    value.append(info)
                    x='$*'.join(map(str,value))
                    data.loc[index,'extracted_sentance']=x
        return (data)

    def mergeDocumentData(self,data1,data2):
        df=pd.merge(data1,data2, on="Department")
        df = df[df.Department != "Total"]
        return df


    